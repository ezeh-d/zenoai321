"""ConversionEngine (#28) -- change a file's format, verified, honestly.

Deterministic converters using libraries ZENO already has (openpyxl, python-docx,
Pillow) for the pairs they cover, plus a headless-LibreOffice route for
Office -> PDF where soffice is installed. The spec's rule is followed: prefer a
real headless converter over an unreliable hand-rolled one, and NEVER claim a
conversion succeeded that didn't -- an unsupported pair or a missing tool is
reported plainly, and every successful write is verified.
"""

from __future__ import annotations

import csv
import io
import os
import shutil
import subprocess
from pathlib import Path
from typing import Any


def _abs(p: str | Path) -> Path:
    return Path(os.path.abspath(os.path.expanduser(str(p))))


def _fmt(path: Path) -> str:
    from reyes_agent.content import format_router as fr
    return fr.detect(path).fmt if path.exists() else path.suffix.lower().lstrip(".")


# --- pure-python converters (installed libs) --------------------------------
def _csv_to_xlsx(src: Path, dest: Path, delim: str = ",") -> None:
    from openpyxl import Workbook
    rows = list(csv.reader(src.read_text(encoding="utf-8", errors="replace")
                           .splitlines(), delimiter=delim))
    wb = Workbook(); ws = wb.active
    for row in rows:
        ws.append(row)
    wb.save(dest)


def _xlsx_to_csv(src: Path, dest: Path) -> None:
    from reyes_agent.content import tables as tbl
    result = tbl.extract_tables(src)
    if not result.get("ok") or not result["tables"]:
        raise ValueError("no sheet/table to export")
    dest.write_text(tbl.to_csv(result["tables"][0]), encoding="utf-8")


def _image_to_pdf(src: Path, dest: Path) -> None:
    from PIL import Image
    with Image.open(src) as im:
        im.convert("RGB").save(dest, "PDF", resolution=150.0)


def _text_to_docx(src: Path, dest: Path) -> None:
    from docx import Document
    doc = Document()
    for line in src.read_text(encoding="utf-8", errors="replace").splitlines() or [""]:
        doc.add_paragraph(line)
    doc.save(dest)


def _docx_to_txt(src: Path, dest: Path) -> None:
    from docx import Document
    doc = Document(src)
    dest.write_text("\n".join(p.text for p in doc.paragraphs), encoding="utf-8")


def _md_to_html(src: Path, dest: Path) -> None:
    import html as _html
    import re
    out: list[str] = ["<!doctype html><meta charset='utf-8'><body>"]
    for line in src.read_text(encoding="utf-8", errors="replace").splitlines():
        s = line.rstrip()
        if not s:
            continue
        h = re.match(r"^(#{1,6})\s+(.*)$", s)
        if h:
            lvl = len(h.group(1))
            out.append(f"<h{lvl}>{_html.escape(h.group(2))}</h{lvl}>")
            continue
        if s.lstrip().startswith(("- ", "* ")):
            out.append(f"<li>{_html.escape(s.lstrip()[2:])}</li>")
            continue
        body = _html.escape(s)
        body = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", body)
        out.append(f"<p>{body}</p>")
    out.append("</body>")
    dest.write_text("\n".join(out), encoding="utf-8")


# (src_fmt, dest_fmt) -> converter
_PURE: dict[tuple[str, str], Any] = {
    ("csv", "xlsx"): _csv_to_xlsx,
    ("tsv", "xlsx"): lambda s, d: _csv_to_xlsx(s, d, "\t"),
    ("xlsx", "csv"): _xlsx_to_csv,
    ("xlsm", "csv"): _xlsx_to_csv,
    ("png", "pdf"): _image_to_pdf, ("jpg", "pdf"): _image_to_pdf,
    ("jpeg", "pdf"): _image_to_pdf, ("bmp", "pdf"): _image_to_pdf,
    ("webp", "pdf"): _image_to_pdf, ("tiff", "pdf"): _image_to_pdf,
    ("txt", "docx"): _text_to_docx, ("md", "docx"): _text_to_docx,
    ("docx", "txt"): _docx_to_txt,
    ("md", "html"): _md_to_html, ("txt", "html"): _md_to_html,
}
# Office -> PDF via headless LibreOffice.
_SOFFICE_PDF = {"docx", "doc", "pptx", "ppt", "xlsx", "xls", "odt", "odp", "ods", "rtf"}


def _soffice() -> str:
    for name in ("soffice", "soffice.exe", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    for guess in (r"C:\Program Files\LibreOffice\program\soffice.exe",
                  r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"):
        if Path(guess).exists():
            return guess
    return ""


def _office_to_pdf(src: Path, dest: Path) -> None:
    exe = _soffice()
    if not exe:
        raise RuntimeError("LibreOffice (soffice) is not installed; cannot "
                           "convert this Office format to PDF reliably")
    out_dir = dest.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run([exe, "--headless", "--convert-to", "pdf", "--outdir",
                    str(out_dir), str(src)], check=True, timeout=120,
                   capture_output=True)
    produced = out_dir / (src.stem + ".pdf")
    if produced != dest and produced.exists():
        os.replace(produced, dest)


def available_conversions() -> dict[str, Any]:
    """What ZENO can actually convert right now, honestly."""
    return {"pure_python": sorted(f"{a}->{b}" for a, b in _PURE),
            "office_to_pdf_via_libreoffice": bool(_soffice()),
            "libreoffice": _soffice() or "not installed"}


def convert(src: str | Path, dest: str | Path) -> dict[str, Any]:
    """Convert src to dest (format from dest's extension), verified. Honest on
    an unsupported pair or a missing tool -- never a fake success."""
    s, d = _abs(src), _abs(dest)
    if not s.exists() or not s.is_file():
        return {"ok": False, "error": f"'{s}' is not a file"}
    sfmt = _fmt(s)
    dfmt = d.suffix.lower().lstrip(".")
    if not dfmt:
        return {"ok": False, "error": "destination has no extension to convert to"}
    if sfmt == dfmt:
        return {"ok": False, "error": f"source and destination are both '{sfmt}'"}

    converter = _PURE.get((sfmt, dfmt))
    route = "pure-python"
    if converter is None and dfmt == "pdf" and sfmt in _SOFFICE_PDF:
        converter, route = _office_to_pdf, "libreoffice"
    if converter is None:
        return {"ok": False, "src_format": sfmt, "dest_format": dfmt,
                "error": f"no converter for {sfmt} -> {dfmt}. "
                         f"Available: {available_conversions()['pure_python']}"}
    try:
        d.parent.mkdir(parents=True, exist_ok=True)
        converter(s, d)
    except Exception as exc:  # noqa: BLE001
        return {"ok": False, "src_format": sfmt, "dest_format": dfmt,
                "route": route, "error": f"{type(exc).__name__}: {exc}"[:200]}

    from reyes_agent.content.save import verify_write
    v = verify_write(d, expect_format=dfmt if dfmt in ("csv", "json", "html") else "")
    return {"ok": v["ok"], "path": str(d), "src_format": sfmt,
            "dest_format": dfmt, "route": route, "verified": v["checks"],
            "error": "" if v["ok"] else v.get("error", "verification failed")}
